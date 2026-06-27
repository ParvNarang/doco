# editor.py
"""
Two FastAPI routers for the collaborative editor feature.

  router       -- prefix /api/editor  -- document CRUD + DOCX export
  agent_router -- prefix /api/agent   -- two-phase SSE agent

Include both in main.py:
    from editor import router as editor_router, agent_router
    app.include_router(editor_router)
    app.include_router(agent_router)
"""

import os
import io
import json
import uuid
import time
import datetime
import logging

from fastapi import APIRouter
from fastapi.responses import JSONResponse, Response, StreamingResponse
from pydantic import BaseModel

import agent as _agent
from config import settings

log = logging.getLogger("editor")

# ─────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────
DOCUMENTS_DIR = "documents"


def _doc_path(doc_id: str) -> str:
    safe = doc_id.replace("/", "").replace("\\", "").replace("..", "")
    return os.path.join(DOCUMENTS_DIR, safe, "document.json")


def _ensure_dir(doc_id: str):
    safe = doc_id.replace("/", "").replace("\\", "").replace("..", "")
    os.makedirs(os.path.join(DOCUMENTS_DIR, safe), exist_ok=True)


# ─────────────────────────────────────────────────────────────
# Router A -- Document CRUD + Export
# ─────────────────────────────────────────────────────────────
router = APIRouter(prefix="/api/editor")


class CreateDocRequest(BaseModel):
    title: str = "Untitled Document"


class SaveDocRequest(BaseModel):
    title: str
    content: list


@router.post("/documents")
async def create_document(req: CreateDocRequest):
    """Create a new empty document. Returns full doc object."""
    doc_id = str(uuid.uuid4())
    now_iso = datetime.datetime.utcnow().isoformat() + "Z"
    _ensure_dir(doc_id)

    doc = {
        "id": doc_id,
        "title": req.title,
        "created_at": now_iso,
        "updated_at": now_iso,
        "content": [{"insert": "\n"}],
    }

    with open(_doc_path(doc_id), "w", encoding="utf-8") as f:
        json.dump(doc, f, ensure_ascii=False, indent=2)

    return JSONResponse(content=doc)


@router.get("/documents")
async def list_documents():
    """List all saved documents, newest-first. Returns metadata only."""
    if not os.path.exists(DOCUMENTS_DIR):
        return JSONResponse(content=[])

    results = []
    for entry in os.scandir(DOCUMENTS_DIR):
        if not entry.is_dir():
            continue
        path = os.path.join(entry.path, "document.json")
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                doc = json.load(f)
            results.append({
                "id": doc["id"],
                "title": doc.get("title", "Untitled"),
                "created_at": doc.get("created_at", ""),
                "updated_at": doc.get("updated_at", ""),
            })
        except Exception:
            continue

    results.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    return JSONResponse(content=results)


@router.get("/documents/{doc_id}")
async def get_document(doc_id: str):
    """Return full document including Quill Delta content."""
    path = _doc_path(doc_id)
    if not os.path.exists(path):
        return JSONResponse(status_code=404, content={"error": "Document not found"})
    with open(path, "r", encoding="utf-8") as f:
        doc = json.load(f)
    return JSONResponse(content=doc)


@router.put("/documents/{doc_id}")
async def save_document(doc_id: str, req: SaveDocRequest):
    """Autosave -- update title and content."""
    if len(json.dumps(req.content)) > 5_000_000:
        return JSONResponse(
            status_code=413,
            content={"error": "Document too large"}
        )

    path = _doc_path(doc_id)
    now_iso = datetime.datetime.utcnow().isoformat() + "Z"

    if not os.path.exists(path):
        _ensure_dir(doc_id)
        doc = {
            "id": doc_id,
            "title": req.title,
            "created_at": now_iso,
            "updated_at": now_iso,
            "content": req.content,
        }
    else:
        with open(path, "r", encoding="utf-8") as f:
            doc = json.load(f)
        doc["title"] = req.title
        doc["content"] = req.content
        doc["updated_at"] = now_iso

    with open(path, "w", encoding="utf-8") as f:
        json.dump(doc, f, ensure_ascii=False, indent=2)

    return JSONResponse(content={"status": "saved", "updated_at": now_iso})


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document and its directory."""
    path = _doc_path(doc_id)
    if not os.path.exists(path):
        return JSONResponse(status_code=404, content={"error": "Document not found"})
    os.remove(path)
    try:
        dir_path = os.path.dirname(path)
        if not os.listdir(dir_path):
            os.rmdir(dir_path)
    except OSError:
        pass
    return JSONResponse(content={"status": "deleted"})


@router.post("/documents/{doc_id}/export/docx")
async def export_docx(doc_id: str):
    """Convert Quill Delta -> .docx and stream it to the browser."""
    path = _doc_path(doc_id)
    if not os.path.exists(path):
        return JSONResponse(status_code=404, content={"error": "Document not found"})

    with open(path, "r", encoding="utf-8") as f:
        doc_data = json.load(f)

    title = doc_data.get("title", "Document")
    delta_ops = doc_data.get("content", [])
    docx_bytes = _delta_to_docx(delta_ops, title)

    safe_title = "".join(c for c in title if c.isalnum() or c in " _-")[:60].strip() or "document"
    filename = f"{safe_title}.docx"

    return Response(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────────────────────
# Quill Delta -> DOCX converter
# ─────────────────────────────────────────────────────────────

def _delta_to_docx(delta_ops: list, title: str) -> bytes:
    """
    Convert a Quill 2 Delta ops array to python-docx bytes.

    Delta encoding rules:
    - Text ops: {insert: "str", attributes?: {bold, italic, underline, ...}}
    - Block attrs (header, list) live on the \\n character that ends each line
    - An op may contain multiple \\n characters -- split on each
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    doc.core_properties.title = title

    run_buffer: list[tuple[str, dict]] = []

    def flush_paragraph(line_attrs: dict):
        la = line_attrs or {}
        header = la.get("header", 0)
        list_type = la.get("list", None)
        blockquote = la.get("blockquote", False)

        if header:
            level = min(int(header), 9)
            p = doc.add_heading("", level=level)
        elif list_type == "bullet":
            p = doc.add_paragraph(style="List Bullet")
        elif list_type == "ordered":
            p = doc.add_paragraph(style="List Number")
        elif blockquote:
            try:
                p = doc.add_paragraph(style="Quote")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
        else:
            p = doc.add_paragraph()

        for text, attrs in run_buffer:
            if not text:
                continue
            run = p.add_run(text)
            if attrs.get("bold"):
                run.bold = True
            if attrs.get("italic"):
                run.italic = True
            if attrs.get("underline"):
                run.underline = True
            if attrs.get("strike"):
                run.font.strike = True
            if attrs.get("code"):
                run.font.name = "Courier New"
                run.font.size = Pt(10)

        run_buffer.clear()

    for op in delta_ops:
        insert = op.get("insert", "")
        attrs = op.get("attributes") or {}

        if not isinstance(insert, str):
            continue

        if "\n" not in insert:
            run_buffer.append((insert, attrs))
            continue

        parts = insert.split("\n")
        for i, part in enumerate(parts):
            if part:
                run_buffer.append((part, attrs))
            if i < len(parts) - 1:
                block_attrs = attrs if (i == len(parts) - 2) else {}
                flush_paragraph(block_attrs)

    if run_buffer:
        flush_paragraph({})

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────
# Router B -- Agent
# ─────────────────────────────────────────────────────────────
agent_router = APIRouter(prefix="/api/agent")


class EvaluateRequest(BaseModel):
    recent_text: str
    document_context: str
    cursor_index: int
    doc_id: str
    model: str = settings.LLM_MODEL


class FeedbackRequest(BaseModel):
    proposal_id: str


@agent_router.post("/evaluate")
async def evaluate(req: EvaluateRequest):
    """
    Two-phase SSE stream:
      Phase 1 -- classify_intervention()
        Emits: {type:"classification", should_intervene, reason, action_type}
        If should_intervene is false -> emits [DONE] and exits.
      Phase 2 -- generate_proposal_stream()
        Emits: {type:"rag_status", result_count}
                {type:"proposal",  proposal_id, action_type, reason,
                                   content, target_text}
                data: [DONE]
    """
    log.info("=" * 60)
    log.info("[evaluate] Request received -- model=%s, doc=%s, cursor=%d",
             req.model, req.doc_id, req.cursor_index)
    log.info("[evaluate] recent_text (%d chars): %.200s...",
             len(req.recent_text), req.recent_text)
    log.info("[evaluate] document_context (%d chars): %.200s...",
             len(req.document_context), req.document_context)

    import main as _main
    log.info("[evaluate] Unloading Surya models to free GPU memory...")
    _main.unload_surya_models()

    req.recent_text = req.recent_text[:800]
    req.document_context = req.document_context[:2000]

    async def _stream():
        t0 = time.time()
        log.info("[stream] Phase 1: Classification starting...")
        cls = _agent.classify_intervention(
            req.recent_text, model_name=req.model
        )

        # If classification itself failed (Ollama down, model missing, etc.),
        # emit an error event so the user sees what went wrong.
        if cls.get("error"):
            elapsed = time.time() - t0
            error_detail = cls["error"]
            log.error("[stream] Classification FAILED after %.1fs: %s", elapsed, error_detail)
            if "Cannot connect" in error_detail or "ConnectError" in error_detail:
                error_detail = "Ollama is not running. Start it with: ollama serve"
            elif "not found" in error_detail.lower() or "status 404" in error_detail:
                error_detail = f"Model '{req.model}' not found. Pull it with: ollama pull {req.model}"
            yield f"data: {json.dumps({'type': 'error', 'error': error_detail})}\n\n"
            yield "data: [DONE]\n\n"
            return

        elapsed = time.time() - t0
        log.info("[stream] Classification done in %.1fs", elapsed)
        yield f"data: {json.dumps({'type': 'classification', **cls})}\n\n"

        if not cls.get("should_intervene", False):
            log.info("[stream] No intervention needed. Reason: %s", cls.get("reason"))
            yield "data: [DONE]\n\n"
            return

        log.info("[stream] Phase 2: RAG + Generation starting (action=%s)...", cls["action_type"])
        t1 = time.time()
        for event in _agent.generate_proposal_stream(
            action_type=cls["action_type"],
            reason=cls["reason"],
            search_query=cls["search_query"],
            document_context=req.document_context,
            recent_text=req.recent_text,
            model_name=req.model,
        ):
            yield event
        elapsed2 = time.time() - t1
        log.info("[stream] Phase 2 done in %.1fs", elapsed2)
        log.info("[stream] Total request time: %.1fs", time.time() - t0)
        log.info("=" * 60)

    return StreamingResponse(_stream(), media_type="text/event-stream")


@agent_router.get("/health")
async def agent_health():
    """
    Check if the agent backend is reachable.
    Tests Ollama connectivity and model availability.
    Returns: {status, ollama_reachable, model, model_available}
    """
    import httpx

    model_name = settings.LLM_MODEL
    result = {
        "status": "ok",
        "ollama_reachable": False,
        "model": model_name,
        "model_available": False,
        "error": None,
    }

    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            # Check Ollama is running
            resp = await client.get("http://localhost:11434/api/tags")
            if resp.status_code == 200:
                result["ollama_reachable"] = True
                data = resp.json()
                models_list = [m.get("name", "") for m in data.get("models", [])]
                # Check exact match or prefix match (ollama lists as "model:tag")
                result["model_available"] = any(
                    model_name == m or model_name in m for m in models_list
                )
            else:
                result["error"] = f"Ollama returned status {resp.status_code}"
    except httpx.ConnectError:
        result["error"] = "Cannot connect to Ollama at localhost:11434"
    except Exception as e:
        result["error"] = str(e)

    if not result["ollama_reachable"]:
        result["status"] = "offline"
    elif not result["model_available"]:
        result["status"] = "no_model"

    return JSONResponse(content=result)


@agent_router.get("/knowledge-status")
async def knowledge_status():
    """
    Check how many documents are available for RAG retrieval.
    Returns: {total_documents, documents: [{uid, filename, chunk_count}]}
    """
    data_dir = "data"
    results = {"total_documents": 0, "documents": []}

    if not os.path.exists(data_dir):
        return JSONResponse(content=results)

    for uid in os.listdir(data_dir):
        uid_dir = os.path.join(data_dir, uid)
        if not os.path.isdir(uid_dir):
            continue

        meta_path = os.path.join(uid_dir, "metadata.json")
        faiss_path = os.path.join(uid_dir, "faiss_index")
        bm25_path = os.path.join(uid_dir, "bm25_index.pkl")

        filename = uid
        if os.path.exists(meta_path):
            try:
                with open(meta_path, "r", encoding="utf-8") as f:
                    filename = json.load(f).get("filename", uid)
            except Exception:
                pass

        has_index = os.path.exists(faiss_path) and os.path.exists(bm25_path)

        results["documents"].append({
            "uid": uid,
            "filename": filename,
            "indexed": has_index,
        })
        results["total_documents"] += 1

    return JSONResponse(content=results)


@agent_router.post("/accept")
async def accept_proposal(req: FeedbackRequest):
    """Log acceptance. Extend later for RLHF data collection."""
    log.info("[agent] ACCEPTED proposal %s", req.proposal_id)
    return JSONResponse(content={"status": "ok"})


@agent_router.post("/reject")
async def reject_proposal(req: FeedbackRequest):
    """Log rejection."""
    log.info("[agent] REJECTED proposal %s", req.proposal_id)
    return JSONResponse(content={"status": "ok"})
